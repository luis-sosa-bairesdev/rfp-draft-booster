"""DOCX export service for RFP Draft Booster.

This module provides functionality to export proposal drafts to .docx format
with automatic formatting.
"""

import logging
import re
from datetime import datetime
from io import BytesIO
from typing import Optional, List, Dict, Any

# python-docx imports
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not available. Install with: pip install python-docx")

from src.models.draft import Draft
from src.models.rfp import RFP

logger = logging.getLogger(__name__)


class DocxExporter:
    """Export drafts to .docx format.
    
    This service handles:
    - Creating formatted .docx documents
    - Markdown to Word conversion
    - Metadata headers
    - Service matches table
    """
    
    def __init__(self):
        """Initialize the DOCX exporter."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx library not available")
    
    def is_available(self) -> bool:
        """Check if DOCX export is available.
        
        Returns:
            True if python-docx is installed, False otherwise
        """
        return DOCX_AVAILABLE
    
    def export_to_docx(
        self,
        draft: Draft,
        rfp: Optional[RFP] = None,
        service_matches: Optional[List[Dict[str, Any]]] = None
    ) -> Optional[bytes]:
        """Export draft to .docx format.
        
        Args:
            draft: Draft to export
            rfp: Optional RFP for metadata
            service_matches: Optional list of service match dictionaries
        
        Returns:
            Bytes of the .docx file, or None if export fails
        """
        if not DOCX_AVAILABLE:
            logger.error("Cannot export to .docx: python-docx not installed")
            return None
        
        try:
            logger.debug("Generating .docx file")
            
            # Create document
            doc = Document()
            
            # Add title
            title = self._generate_doc_title(draft, rfp)
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add metadata
            if rfp:
                self._add_metadata(doc, rfp)
            
            # Add service matches table
            if service_matches:
                self._add_service_matches_table(doc, service_matches)
            
            # Add main content
            self._add_markdown_content(doc, draft.content)
            
            # Save to bytes
            docx_bytes = BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            logger.info("Successfully generated .docx file")
            return docx_bytes.getvalue()
        
        except Exception as e:
            logger.error(f"Failed to generate .docx: {e}", exc_info=True)
            return None
    
    def _generate_doc_title(self, draft: Draft, rfp: Optional[RFP]) -> str:
        """Generate document title.
        
        Args:
            draft: Draft object
            rfp: Optional RFP object
        
        Returns:
            Document title string
        """
        if rfp and hasattr(rfp, 'id'):
            rfp_id = rfp.id[:8] if len(rfp.id) > 8 else rfp.id
            return f"RFP Draft - {rfp_id} - {datetime.now().strftime('%Y-%m-%d')}"
        return f"Draft - {datetime.now().strftime('%Y-%m-%d')}"
    
    def _add_metadata(self, doc: Document, rfp: RFP):
        """Add RFP metadata section to document.
        
        Args:
            doc: Document object
            rfp: RFP object
        """
        doc.add_heading("RFP Information", level=2)
        
        metadata_table = doc.add_table(rows=0, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        # Add metadata rows
        metadata = [
            ("RFP ID", getattr(rfp, 'id', 'N/A')[:8] if hasattr(rfp, 'id') else 'N/A'),
            ("Organization", getattr(rfp, 'organization', 'N/A')),
            ("Submission Date", getattr(rfp, 'submission_deadline', 'N/A')),
            ("Generated", datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for key, value in metadata:
            row = metadata_table.add_row()
            row.cells[0].text = key
            row.cells[1].text = str(value)
        
        doc.add_paragraph()  # Spacing
    
    def _add_service_matches_table(self, doc: Document, service_matches: List[Dict[str, Any]]):
        """Add service matches table to document.
        
        Args:
            doc: Document object
            service_matches: List of service match dictionaries
        """
        doc.add_heading("Recommended Services", level=2)
        
        if not service_matches:
            doc.add_paragraph("No service matches available.")
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Requirement"
        header_cells[1].text = "Service"
        header_cells[2].text = "Match %"
        
        # Data rows
        for match in service_matches:
            row = table.add_row()
            row.cells[0].text = str(match.get('requirement_desc', ''))[:100]
            row.cells[1].text = str(match.get('service_name', 'N/A'))
            row.cells[2].text = f"{match.get('match_percentage', 0):.1f}%"
        
        doc.add_paragraph()  # Spacing
    
    def _add_markdown_content(self, doc: Document, content: str):
        """Add markdown content to document with basic formatting.
        
        Args:
            doc: Document object
            content: Markdown content
        """
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Skip empty lines
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                para = doc.add_paragraph(text, style='List Number')
            
            # Regular paragraph with inline formatting
            else:
                para = doc.add_paragraph()
                self._add_inline_formatting(para, line)
            
            i += 1
    
    def _add_inline_formatting(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic) to paragraph.
        
        Args:
            paragraph: Paragraph object
            text: Text with markdown formatting
        """
        # Simple regex-based inline formatting
        # Handle **bold**, *italic*, and `code`
        
        # Replace **bold**
        pattern = r'\*\*(.+?)\*\*'
        parts = re.split(pattern, text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text, check for *italic*
                italic_pattern = r'\*(.+?)\*'
                italic_parts = re.split(italic_pattern, part)
                for j, italic_part in enumerate(italic_parts):
                    if j % 2 == 0:
                        # Regular text, check for `code`
                        code_pattern = r'`(.+?)`'
                        code_parts = re.split(code_pattern, italic_part)
                        for k, code_part in enumerate(code_parts):
                            if k % 2 == 0:
                                if code_part:
                                    paragraph.add_run(code_part)
                            else:
                                run = paragraph.add_run(code_part)
                                run.font.name = 'Courier New'
                    else:
                        run = paragraph.add_run(italic_part)
                        run.italic = True
            else:
                run = paragraph.add_run(part)
                run.bold = True

